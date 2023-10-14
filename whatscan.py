# coding = utf-8
# #!/usr/bin/python3
# pip install selenium python-docx translate docxcompose scikit-learn beautifulsoup4

import csv
import os
import re
import time
import html
import docx     
import datetime    
import ctypes
import socket
import requests
import traceback
import common.tf_idf as tf_idf
from queue import Queue
from threading import Thread
from urllib.parse import urlparse
from urllib.parse import quote
from threading import active_count
from configparser import ConfigParser
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from translate import Translator   
requests.packages.urllib3.disable_warnings()

print('\nPlease install chrome browser and download suitable chromedriver.exe(put it on current directory) before using this tools.( https://chromedriver.chromium.org/downloads)\n')

try:
    ini_obj = ConfigParser()
    ini_obj.read('config.ini', encoding='utf-8')
    browser_thread = int(ini_obj.get('set','browser_thread'))
    input_file = ini_obj.get('set','input_file')
    read_index = int(ini_obj.get('set','read_index'))
    q_input_length = ini_obj.get('set','q_input_length')
    q_output_length = ini_obj.get('set','q_output_length')
    use_proxy = int(ini_obj.get('set','use_proxy'))
    proxy_server = ini_obj.get('set','proxy_server')
    proxy_port = ini_obj.get('set','proxy_port')
    google_tran_api_timeout = int(ini_obj.get('set','google_tran_api_timeout'))
    page_load_timeout = int(ini_obj.get('set','set_page_load_timeout'))
    need_tanslate = int(ini_obj.get('set','need_tanslate'))
    need_word_freq = int(ini_obj.get('set','need_word_freq'))
    translate_using_proxy = int(ini_obj.get('set','translate_using_proxy'))
except Exception as e:
    print("parse config.ini error: ", e)
    exit(1)
    
# 加载DLL
kscan_lib = ctypes.CDLL('./kscan.dll')
# 定义参数类型
kscan_lib.Search_web.argtypes = [
    ctypes.c_char_p,  # Protocol
    ctypes.c_char_p,  # Port
    ctypes.c_char_p,  # Header
    ctypes.c_char_p,  # Body
    ctypes.c_char_p,  # Response
    ctypes.c_char_p,  # Cert
    ctypes.c_char_p,  # Title
    ctypes.c_char_p,  # Hash
    ctypes.c_char_p,  # Icon
    ctypes.c_char_p   # ICP
]

# 定义返回值类型
kscan_lib.Search_web.restype = ctypes.c_char_p

q_in = Queue(maxsize=int(q_input_length))
q_out = Queue(maxsize=int(q_output_length))
q_csv = Queue(maxsize=100)
seconds = time.time()
time_str = time.strftime("%Y-%m-%d-%H-%M-%S", time.localtime(seconds))
output_path = f"{os.path.dirname(os.path.abspath(__file__))}/output/{time_str}"
tmp_file_path = f"{os.path.dirname(os.path.abspath(__file__))}/output/{time_str}/tmp"
log_path = f"{output_path}"
translator = Translator(to_lang="chinese")  # set target language to CN
proxies = {'http':f'http://{proxy_server}:{proxy_port}', 'https':f'http://{proxy_server}:{proxy_port}'}  if use_proxy else {}
trans_proxy = proxies if translate_using_proxy  else  {}
    
is_exit = False
driver_list = []
GOOGLE_TRANSLATE_URL = 'https://translate.google.com/m?q=%s&tl=%s&sl=%s'


def google_translate(text, to_language="auto", text_language="auto"):
    text = quote(text)
    url = GOOGLE_TRANSLATE_URL % (text,to_language,text_language)
    response = requests.get(url, proxies=trans_proxy, timeout=google_tran_api_timeout, verify=False)
    data = response.text
    expr = r'(?s)class="(?:t0|result-container)">(.*?)<'
    result = re.findall(expr, data)
    if (len(result) == 0):
        return ""
    return html.unescape(result[0])

def isip(str):
    p = re.compile(r'^((25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(25[0-5]|2[0-4]\d|[01]?\d\d?(:[0-9]+)?/.*?)$') # fitter 'ip' or 'ip:port' or 'ip:port/path'
    if p.match(str) :
        return True
    else:
        return False

def iter_count(file_name):
    try:
        from itertools import (takewhile, repeat)
        buffer = 1024 * 1024
        with open('input/' + file_name) as f:
            buf_gen = takewhile(lambda x: x, (f.read(buffer) for _ in repeat(None)))
            return sum(buf.count('\n') for buf in buf_gen)
    except :
        return None

def kscan_verify(Protocol, Port, Header, Body, Response, Title=b'', Hash=b'', Icon=b'', IPC=b'', Cert=b''):
    result_ptr = kscan_lib.Search_web(Protocol, Port, Header, Body, Response, Cert, Title, Hash, Icon, IPC)

    # 将C字符串指针转换为Python字符串
    result_str = ctypes.string_at(result_ptr).decode('utf-8', errors='ignore').strip()

    if result_str:
        return result_str
    else:
        return "unknow"



def producer(q_in):
    total_line_count = iter_count(input_file)
    current_index = 0
    with open('input/' + input_file, 'r') as fp:
        if read_index-1 != 0:
            for x in range(0, read_index-1):
                fp.readline()
                current_index += 1
        while True:
            if q_in.full() is False:
                line = fp.readline()
                current_index += 1
                if line:
                    _line = line.strip()
                    if len(_line) > 0 :
                        q_in.put(_line)
                        continue
                else:
                    break
            else:
                time.sleep(1)
                try:
                    if total_line_count == 0: total_line_count = 1
                    current_process = str( (current_index / total_line_count) * 100)[:4] + '%'
                    if current_process[:3] == '100': current_process = current_process[:3] + '%'
                    print(f"current progress: {current_process} , {current_index}/{total_line_count}")
                    with open(log_path + '/progress.txt', 'w') as f_progress:
                        f_progress.write(f"current progress: {current_process} , {current_index}/{total_line_count}")
                except Exception as e:
                    e_info = f"exception, {e.__traceback__.tb_frame.f_globals['__file__']}, line: {e.__traceback__.tb_lineno}\n{str(e)}\n"
                    print(e_info)




def exploit(q_in, q_out):
    chrome_driver_path = r'chromedriver.exe'
    chrome_options = Options()
    chrome_options.add_argument('--silent')
    chrome_options.add_argument('log-level=3')
    chrome_options.add_experimental_option("excludeSwitches",['enable-automation','enable-logging'])
    chrome_options.add_argument('--lang=en-US') 
    chrome_options.add_argument('--ignore-certificate-errors')  # ignore ssl certificate errors
    chrome_options.add_argument('--disable-software-rasterizer')
    chrome_options.add_argument('window-size=1440x900')
    chrome_options.add_argument('--headless')     # don't display windows
    chrome_options.add_argument('--disable-gpu')  # don't display windows
    chrome_options.add_argument('--disable-extensions')  
    chrome_options.add_argument("disable-cache")
    if use_proxy:
        chrome_options.add_argument(f'--proxy-server={proxy_server}:{proxy_port}')    # set http proxy
    '''
        #set socks proxy
        proxy = Proxy()
        proxy.proxyType = ProxyType.MANUAL
        proxy.autodetect = False
        proxy.httpProxy = proxy.sslProxy = proxy.socksProxy = "127.0.0.1:8119"
        chrome_options.Proxy = proxy
    '''


    service = Service(executable_path='chromedriver.exe')
    try:
        driver = webdriver.Chrome(options=chrome_options, service=service)
    except:
        driver = webdriver.Chrome(options=chrome_options, executable_path=chrome_driver_path)
    
    driver.set_page_load_timeout(page_load_timeout)
    driver_list.append(driver)

    while True:
        doc = None
        is_cdn = False
        try:
            target = q_in.get(timeout=5)
        except :
            break
        
        try:
            try:
                driver.get(target, )
                doc = docx.Document('Normal.docx')
            except :
                continue

            url = urlparse(target)
            hostAndPort = url.netloc
            ip_addr = None
            ptr_domain = None
            target_port = None
            try:
                hostname, target_port = hostAndPort.split(":")
                ip_addr = hostname
                hostAndPort = hostAndPort.replace(':','_')
            except Exception as e:
                hostname = hostAndPort
                if url.scheme == 'http':
                    target_port = '80'
                else:
                    target_port = '443'
            ts = str(datetime.datetime.now().timestamp())
            img_name = f"output/{time_str}/png/{hostAndPort}.{ts}.png"
            try:
                driver.get_screenshot_as_file(img_name)
            except Exception as e:
                e_info = f"exception, {e.__traceback__.tb_frame.f_globals['__file__']}, line: {e.__traceback__.tb_lineno}\n{str(e)}\n"
                print(e_info)
                img_name = f"output/{time_str}/png/{ts}.png"
                driver.get_screenshot_as_file(img_name)
            title = driver.title
            
            
            if isip(hostname):
                pass
            else:
                try:
                    ips = socket.gethostbyname_ex(hostname)
                    if len(ips[2]) >= 3 : 
                        # print('maybe cdn')
                        is_cdn = True
                    ip_addr = ips[2][0]
                except socket.error:
                    pass

            tf_idf_list = None
            
            if need_word_freq:
                tf_idf_list = tf_idf.get_tf_idf_sort_list(driver.page_source)
            
            _be_transed = 'None'
            try:
                
                doc.add_heading(title, 1) 
                doc.add_paragraph(target)
                try:
                    _be_transed = google_translate(title, to_language='zh-CN')
                    doc.add_paragraph('标题翻译: ' + _be_transed)
                except:
                    pass
       
                if is_cdn:
                    doc.add_paragraph('IS CDN: More then 3 IP, maybe')
                else:
                    doc.add_paragraph('IS CDN: No')
                
                try:
                    ptr_domain = socket.gethostbyaddr(ip_addr)
                    doc.add_paragraph(f'IP反析域名: {ptr_domain}')
                    doc.add_paragraph('')
                except:
                    pass
                
                if tf_idf_list:
                    doc.add_paragraph('词频:')
                    combine_word = ''
                    
                    for word, k in tf_idf_list:
                        combine_word += word 
                        combine_word += '|'
                    
                    _word_transed = None
                    if need_tanslate:
                        try:
                            _word_transed = google_translate(combine_word, to_language='zh-CN')
                        except Exception as e:
                            print(e)
                        
                    if _word_transed:
                        _words = _word_transed.split("|")
                        for i in range(0, len(tf_idf_list)):
                            word = tf_idf_list[i][0]
                            word_fq = str(tf_idf_list[i][1])[:3]
                            try:
                                _current_translated_word = _words[i]
                            except:
                                _current_translated_word = word
                            doc.add_paragraph(f'\t{word} ({_current_translated_word}): {word_fq}')
                    else:
                        for word, k in tf_idf_list:
                            _word_transed = ''
                            k = str(k)[:3]
                            doc.add_paragraph(f'{word} ({_word_transed}): {k}')
                        
                #kscan_verify(Protocol, Port, Header, Body, Response, Title=b'', Hash=b'', Icon=b'', IPC=b'', Cert=b'')
                try:
                    _req = 'None'
                    Header = b''
                    _content = b''
                    Body = driver.page_source.encode('utf-8', errors='ignore')
                    Response = Body
                    try:
                        _req = requests.head(target, verify=False, timeout=9)
                    except:
                        traceback.print_exc()
                    if _req != 'None':
                        _content = _req.content
                        Header = "\r\n".join( [ i+': '+_req.headers[i] for i in _req.headers ] ).encode('utf-8', errors='ignore')
                        Response = Header + b'\r\n\r\n' + Response
                        if _req.headers.get('server'):
                            doc.add_paragraph("")
                            doc.add_paragraph("Server:\t" + _req.headers.get('server'))
                        elif _req.headers.get('Server'):
                            doc.add_paragraph("")
                            doc.add_paragraph("Server:\t" + _req.headers.get('Server'))
                        
                    kscan_info = kscan_verify(url.scheme.encode('utf-8', errors='ignore'), target_port.encode('utf-8', errors='ignore'), Header, Body, Response, title.encode('utf-8', errors='ignore'), Hash=b'', Icon=b'', IPC=b'', Cert=b'')
                    if kscan_info:
                        doc.add_paragraph("产品:\t" + kscan_info)
                except:
                    traceback.print_exc()
                    
                doc.add_paragraph("")

            except Exception as e:
                e_info = f"exception, {e.__traceback__.tb_frame.f_globals['__file__']}, line: {e.__traceback__.tb_lineno}\n{str(e)}\n"
                print(e_info)
            finally:
                q_csv.put([target, title, _be_transed])
            
            try:
                doc.add_picture(img_name,
                                width=docx.shared.Cm(15.5),
                                height=docx.shared.Cm(10))
            except Exception as e:
                e_info = f"exception, {e.__traceback__.tb_frame.f_globals['__file__']}, line: {e.__traceback__.tb_lineno}\n{str(e)}\n"
                # print(e_info)
            
            doc.save(tmp_file_path + '/' + str(datetime.datetime.now().timestamp()) + '.docx')
            
        except Exception as e:
            e_info = f"exception, {e.__traceback__.tb_frame.f_globals['__file__']}, line: {e.__traceback__.tb_lineno}\n{str(e)}\n"
            print(e_info)
        q_out.put({'success.txt': target})
    
    driver.quit()
    

def consumer_log(q_out):   
    while True:
        output = q_out.get()
        try:
            for _path in output.keys():
                with open(f"{log_path}/{_path}", 'a+') as fp:
                    data = output[_path]
                    print(' ',data, 'ok')
                    fp.write(data)
                    fp.write('\n')
        except Exception as e:
            print(f'[debug] log error, {str(e)} \n')


def combine_all_docx(files_list):
    number_of_sections = len(files_list)
    master = docx.Document('Normal.docx')
    composer = Composer(master)
 
    for i in range(0, number_of_sections):
        try:
            doc_temp = docx.Document((files_list[i]))
            composer.append(doc_temp)
        except Exception as e:
            e_info = f"exception, {e.__traceback__.tb_frame.f_globals['__file__']}, line: {e.__traceback__.tb_lineno}\n{str(e)}\n"
            # print(e_info)
    composer.save(f'output/{time_str}/{time_str}.docx')


def write_csv():
    while True:
        line = q_csv.get()
        fp_csv = open(f'{output_path}/output.xlsx', "a+", encoding='utf_8_sig', errors='ignore', newline='')
        csv_writer = csv.writer(fp_csv, dialect="excel", )
        csv_writer.writerow(line)
        fp_csv.close()



if __name__ == '__main__':
    start_time = time.time() 
    if os.path.exists(output_path) is False:
        os.makedirs(output_path)
    if os.path.exists('input') is False:
        os.makedirs('input')
        print('put input.txt to ./input/')
    if os.path.exists(f"output/{time_str}/png") is False:
        os.makedirs(f"output/{time_str}/png")
    if os.path.exists(tmp_file_path) is False:
        os.makedirs(tmp_file_path)
        
    fp_csv = open(f'{output_path}/output.xlsx', "a+", encoding='utf_8_sig', errors='ignore', newline='')
    csv_writer = csv.writer(fp_csv, dialect="excel", )
    csv_writer.writerow(['URL', 'Title', 'Title_trans'])
    fp_csv.close()

        
    producer_thread = Thread(target=producer, args=(q_in,), name='producer', daemon=True)
    producer_thread.start()
    
    csv_thread = Thread(target=write_csv, args=(), name='csv_writer', daemon=True)
    csv_thread.start()
    
    

    for i in range(1, browser_thread):
        exp_thread = Thread(target=exploit, args=(q_in, q_out), daemon=True)
        exp_thread.start()

    t_log = Thread(target=consumer_log, args=(q_out,), daemon=True)
    t_log.start()


    while active_count() > 3:
        print('[*]thread alive : ' + str(active_count()) )
        time.sleep(3)

    
    from glob import glob
    from docxcompose.composer import Composer
    path_list = glob(tmp_file_path + '/*.docx')
    combine_all_docx(path_list)

    end_time = time.time()   
    run_time = end_time - start_time 
    print("cost time: %s s" % run_time)
    print('done')
    